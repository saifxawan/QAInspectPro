from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_sqa_assignment():
    doc = Document()

    # --- Header Style ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # --- Title Section ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('SQA Assignment: Role of SQA Engineer in SDLC & SRS Development\n')
    run.bold = True
    font_title = run.font
    font_title.size = Pt(20)
    font_title.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f'Course Code: CLO2: C3 Apply, GA3\nProject Name: QAInspect Pro\nSubmission Date: April 2, 2026')
    run.font.size = Pt(12)
    run.italic = True

    doc.add_section()

    # --- Part 1: Project Overview ---
    doc.add_heading('Part 1: Project Overview', level=1)
    
    doc.add_heading('Project Title', level=2)
    doc.add_paragraph('QAInspect Pro - Enterprise Intelligence & Test Management Suite', style='List Bullet')

    doc.add_heading('Purpose of the System', level=2)
    p = doc.add_paragraph('QAInspect Pro is designed to bridge the gap between automated scanning and formal test management. It provides a centralized hub for SQA Engineers to perform security audits, performance benchmarking, and manage a massive repository of over 1,000 industry-standard test cases. The system aims to replace manual spreadsheets and fragmented tools with a unified, high-performance platform.')

    doc.add_heading('Target Users', level=2)
    users = [
        "SQA Engineers: To execute automated scans and manage test repositories.",
        "Security Analysts: To review SSL/TLS and security header status.",
        "DevOps/SREs: To monitor performance metrics like TTLB and payload efficiency.",
        "Project Managers: To view high-level executive reports and quality trends."
    ]
    for user in users:
        doc.add_paragraph(user, style='List Bullet')

    doc.add_heading('Key Features', level=2)
    features = [
        "Automated Intelligence Engine: Real-time URL analysis for security and network health.",
        "Security Audit Pro: Automated checks for HSTS, CSP, X-Frame-Options, and SSL expiry.",
        "Performance Matrix: Detailed breakdown of response times, payload sizes, and efficiency scores.",
        "1000+ Industry-Standard Test Cases: Pre-built functional, security, and performance test templates.",
        "Executive Reporting Dashboard: Dynamic glassmorphic UI with CSV export and historical tracking."
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')

    # --- Part 2: Role of SQA Engineer in SDLC ---
    doc.add_page_break()
    doc.add_heading('Part 2: Role of SQA Engineer in SDLC', level=1)
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'SDLC Phase'
    hdr_cells[1].text = 'SQA Engineer Activities'
    hdr_cells[2].text = 'Artifacts/Deliverables Produced'
    hdr_cells[3].text = 'Contribution to Quality'
    hdr_cells[4].text = 'Preparation for Next Phase'

    sdlc_data = [
        ("Requirement Analysis", "Reviewing SRS for clarity, completeness, and testability. Identifying potential risks.", "Requirement Traceability Matrix (RTM), QA Strategy Document.", "Ensures building the right system by preventing 'garbage-in.'", "Identifying tools and resources needed for the system."),
        ("Design Phase", "Reviewing system architecture and database design (e.g., SQLAlchemy mappings).", "Test Plan, High-Level Test Design.", "Detects architectural flaws (logic errors) before code is written.", "Mapping test cases to specific design modules."),
        ("Development", "White-box testing, Code reviews, and setting up Test Data (e.g., seed_data.py).", "Unit Test Results, Static Analysis Reports.", "Prevents bugs from entering the main codebase; verifies logic early.", "Finalizing test scripts for integration and system testing."),
        ("Testing Phase", "Execution of Functional, Security, and Performance tests on the build.", "Bug Reports (Jira), Test Execution Reports.", "Verifies the system meets requirements and is stable for production.", "Creating a 'Go/No-Go' readiness report for deployment."),
        ("Deployment/Maint.", "Smoke testing on production and verifying fixes for reported field bugs.", "Release Notes, Final Quality Sign-off.", "Ensures zero regression and validates the system in the real user environment.", "Updating regression suites for future cycles.")
    ]

    for phase, activities, artifacts, contribution, prep in sdlc_data:
        row_cells = table.add_row().cells
        row_cells[0].text = phase
        row_cells[1].text = activities
        row_cells[2].text = artifacts
        row_cells[3].text = contribution
        row_cells[4].text = prep

    # --- Part 3: SRS ---
    doc.add_page_break()
    doc.add_heading('Part 3: Software Requirements Specification (SRS)', level=1)
    
    doc.add_heading('1. Introduction', level=2)
    doc.add_heading('1.1 Purpose', level=3)
    doc.add_paragraph('The purpose of this document is to define the functional and non-functional requirements for QAInspect Pro v1.0. It provides a baseline for development and testing.')
    
    doc.add_heading('1.2 Scope', level=3)
    doc.add_paragraph('QAInspect Pro is a web-based intelligence scanner and test management system. It integrates a FastAPI backend with a React frontend to provide real-time quality insights.')
    
    doc.add_heading('1.3 Definitions', level=3)
    defs = [
        "TTLB: Time To Last Byte (Total response time).",
        "HSTS: HTTP Strict Transport Security.",
        "Glassmorphism: A design style characterized by background blur and translucent layers."
    ]
    for d in defs:
        doc.add_paragraph(d, style='List Bullet')

    doc.add_heading('2. Overall Description', level=2)
    doc.add_heading('2.1 Product Perspective', level=3)
    doc.add_paragraph('QAInspect Pro is a standalone intelligence suite that can be integrated into CI/CD pipelines via API. It utilizes a PostgreSQL database for historical data persistence.')
    
    doc.add_heading('2.2 User Characteristics', level=3)
    doc.add_paragraph('Users are expected to have a basic understanding of web security (HTTPS/Headers) and SQA terminologies.')
    
    doc.add_heading('2.3 Assumptions and Constraints', level=3)
    doc.add_paragraph('Assumptions: Target URLs are publicly accessible via the Internet.', style='List Bullet')
    doc.add_paragraph('Constraints: Scans are limited to HTTP/HTTPS protocols only.', style='List Bullet')

    doc.add_heading('3. Functional Requirements', level=2)
    frs = [
        "FR1: Target Scanning: The system shall allow users to input a URL and initiate a multi-threaded security and performance scan.",
        "FR2: Security Validation: The system shall check for the presence of mandatory security headers (CSP, HSTS, X-Content-Type).",
        "FR3: Test Case Repository: The system shall store and display 1,000+ pre-defined test cases categorized by functional domains.",
        "FR4: Report Export: Users shall be able to export scan results and test case statuses in CSV format."
    ]
    for fr in frs:
        doc.add_paragraph(fr, style='List Number')

    doc.add_heading('4. Non-Functional Requirements', level=2)
    nfrs = [
        "NFR1: Performance: The intelligence engine must complete a basic URL scan (headers + latency) in under 5 seconds.",
        "NFR2: UI/UX: The dashboard must adhere to the Glassmorphism design system with a minimum 90% accessibility score (Lighthouse).",
        "NFR3: Scalability: The system shall support concurrent scanning of 5 different URLs without performance degradation."
    ]
    for nfr in nfrs:
        doc.add_paragraph(nfr, style='List Number')

    # --- Part 4: Making Requirements Testable ---
    doc.add_page_break()
    doc.add_heading('Part 4: Making Requirements Testable & Measurable', level=1)
    
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    hdr2_cells = table2.rows[0].cells
    hdr2_cells[0].text = 'Original (Raw) Requirement'
    hdr2_cells[1].text = 'Improved (Testable) Requirement'
    hdr2_cells[2].text = 'Measurable Criteria'

    testable_data = [
        ('"The system should be very fast when scanning."', '"The backend scanner must return the complete header and latency analysis for a standard URL within a specific time frame."', "Response Time < 5.0 seconds for 95% of scanned URLs on a 100Mbps connection."),
        ('"The dashboard should handle many test cases."', '"The frontend must render a list of 1,000 test cases with smooth scrolling and filter responsiveness."', "Frame Rate >= 60 FPS during scrolling; filtering must update in < 200ms."),
        ('"Users should be able to see security risks."', '"The system shall flag any missing security headers (HSTS, CSP, XFO) as CRITICAL in the report."', "100% Accuracy in header detection validated against OWASP ZAP benchmarks."),
        ('"The system should be easy to use."', '"Users must be able to navigate from the home screen to a full scan report in no more than 3 clicks."', "Success rate of 98% in usability testing for first-time SQA users.")
    ]

    for raw, improved, criteria in testable_data:
        row_cells = table2.add_row().cells
        row_cells[0].text = raw
        row_cells[1].text = improved
        row_cells[2].text = criteria

    # Save
    file_path = 's:\\SAIFI\\PROJECTS\\SQA_Project\\SQA_Assignment_QAInspectPro.docx'
    doc.save(file_path)
    print(f"File saved successfully at: {file_path}")

if __name__ == "__main__":
    create_sqa_assignment()
